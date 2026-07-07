"""Word report generation from engine results + assessment.

Deterministic: every number comes from the results object, every verdict from
assess.py. The AI report-writer (needs API credentials) will later add narrative
sections; this module already produces a complete, citable results report.
"""
import datetime

from docx import Document
from docx.shared import Pt


def _table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = "" if v is None else str(v)
    return t


def _fmt(v):
    return f"{v:.3f}" if isinstance(v, (int, float)) else (v or "")


def build_report(dataset_meta: dict, request: dict, results: dict, assessment: dict,
                 interpretation: dict | None = None, mga: dict | None = None) -> Document:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    doc.add_heading("PLS-SEM Analysis Report", level=0)
    section_no = [0]

    def add_section(title):
        section_no[0] += 1
        doc.add_heading(f"{section_no[0]}. {title}", level=1)
    meta = results["meta"]
    doc.add_paragraph(
        f"Generated {datetime.date.today().isoformat()} · dataset: {dataset_meta.get('filename')} "
        f"(n = {meta['n']}) · engine: {meta['engine']} · bootstrap: {meta['nboot']:,} resamples, "
        f"fixed seed {meta['seed']}."
    )

    # --- Method ---------------------------------------------------------------
    add_section("Method")
    n_constructs = len(request["constructs"])
    modes = {}
    for c in request["constructs"]:
        modes.setdefault(c["measurement"], []).append(c["name"])
    mode_text = "; ".join(
        f"{m.replace('_', '-')}: {', '.join(names)}" for m, names in modes.items()
    )
    missing_text = (
        f"Values coded {request['missing_value']} were treated as missing and handled by "
        f"mean replacement." if request.get("missing_value") else
        "No missing-value code was specified."
    )
    doc.add_paragraph(
        f"Partial least squares structural equation modeling (PLS-SEM) was applied to a model "
        f"with {n_constructs} constructs ({mode_text}) and {len(request['paths'])} structural "
        f"paths. {missing_text} Significance was assessed with {meta['nboot']:,} bootstrap "
        f"resamples using percentile confidence intervals."
    )
    interactions = request.get("interactions") or []
    if interactions:
        terms = ", ".join(f"{i['iv']} × {i['moderator']}" for i in interactions)
        doc.add_paragraph(
            f"Moderation was modeled with the two-stage approach (interaction term(s): "
            f"{terms}; Becker et al., 2018)."
        )
    hocs = [c for c in request["constructs"]
            if c["measurement"].startswith("higher_order")]
    if hocs:
        hoc_text = "; ".join(f"{c['name']} ({', '.join(c['dimensions'])})" for c in hocs)
        doc.add_paragraph(
            f"Higher-order construct(s) were estimated with the two-stage (disjoint) "
            f"approach: {hoc_text} (Sarstedt et al., 2019)."
        )
    audit = dataset_meta.get("audit") or {}
    if audit.get("findings"):
        doc.add_paragraph(
            f"Data screening produced {audit['n_findings']} finding(s) "
            f"({', '.join(sorted({f['check'] for f in audit['findings']}))}); "
            f"details are stored with the dataset audit log."
        )

    # --- Measurement model ------------------------------------------------------
    add_section("Measurement Model Assessment")
    rel = [m for m in assessment["measurement_model"] if m["family"] == "internal_consistency"]
    ave = {m["construct"]: m for m in assessment["measurement_model"] if m["family"] == "convergent_validity"}
    by_construct = {}
    for m in rel:
        by_construct.setdefault(m["construct"], {})[m["metric"]] = m
    rows = []
    for name, metrics in by_construct.items():
        a = metrics.get("Cronbach's alpha", {})
        rA = metrics.get("rho_A", {})
        rC = metrics.get("composite reliability rho_C", {})
        av = ave.get(name, {})
        verdicts = {m.get("verdict") for m in (a, rA, rC, av) if m}
        rows.append([name, _fmt(a.get("value")), _fmt(rA.get("value")), _fmt(rC.get("value")),
                     _fmt(av.get("value")), "OK" if verdicts <= {"pass"} else "REVIEW"])
    if rows:
        doc.add_paragraph("Internal consistency reliability and convergent validity "
                          "(criteria: 0.70–0.95 for reliability; AVE ≥ 0.50; Hair et al., 2022):")
        _table(doc, ["Construct", "α", "ρA", "ρC", "AVE", "Verdict"], rows)

    htmt = [m for m in assessment["measurement_model"] if m["family"] == "discriminant_validity"]
    if htmt:
        doc.add_paragraph()
        doc.add_paragraph("Discriminant validity — HTMT (criterion < 0.85; Henseler et al., 2015):")
        _table(doc, ["Construct pair", "HTMT", "Verdict"],
               [[m["construct"], _fmt(m["value"]), m["verdict"].upper()] for m in htmt])

    low = [m for m in assessment["measurement_model"]
           if m["family"] == "indicator_reliability" and m["verdict"] == "fail"]
    doc.add_paragraph()
    doc.add_paragraph(
        "All reflective indicator loadings meet the 0.708 criterion." if not low else
        f"{len(low)} indicator(s) fall below the 0.708 loading criterion: "
        + ", ".join(f"{m['item']} ({_fmt(m['value'])})" for m in low) + "."
    )

    formative = [m for m in assessment["measurement_model"]
                 if m["family"] == "formative_indicator_validity"]
    if formative:
        doc.add_paragraph()
        doc.add_paragraph(
            "Formative indicator validity — outer weight significance "
            "(95% bootstrap CI; n.s. weights retained when the loading is ≥ 0.5; "
            "Hair et al., 2022):"
        )
        _table(doc, ["Construct", "Indicator", "Weight", "95% CI", "Verdict"],
               [[m["construct"], m["item"], _fmt(m["value"]),
                 f"[{m['ci_95'][0]:.3f}; {m['ci_95'][1]:.3f}]",
                 m["verdict"].upper()] for m in formative])

    # --- Structural model -------------------------------------------------------
    add_section("Structural Model Assessment")
    vif_bad = [s for s in assessment["structural_model"]
               if s["family"] == "collinearity" and s["verdict"] != "pass"]
    doc.add_paragraph(
        "Collinearity: all inner VIF values are below 3." if not vif_bad else
        "Collinearity: " + "; ".join(f"{s['construct']} VIF = {_fmt(s['value'])} ({s['verdict']})"
                                     for s in vif_bad) + "."
    )
    r2 = [s for s in assessment["structural_model"] if s["family"] == "explanatory_power"]
    if r2:
        _table(doc, ["Endogenous construct", "R²", "Interpretation"],
               [[s["construct"], _fmt(s["value"]), s["verdict"]] for s in r2])

    doc.add_paragraph()
    doc.add_paragraph("Path coefficients with bootstrap significance "
                      "(95% percentile CIs; verdict: CI excludes zero):")
    _table(doc, ["Hypothesis", "Path", "β", "t", "95% CI", "Verdict"],
           [[h["hypothesis"], h["path"], _fmt(h["estimate"]), _fmt(h["t_value"]),
             f"[{h['ci_95'][0]:.3f}; {h['ci_95'][1]:.3f}]" if h["ci_95"] else "",
             h["verdict"]] for h in assessment["hypotheses"]])

    mediation = assessment.get("mediation") or []
    if mediation:
        doc.add_paragraph()
        doc.add_paragraph(
            "Mediation — specific indirect effects with 95% bootstrap percentile CIs; "
            "classification follows Zhao, Lynch & Chen (2010) / Hair et al. (2022):"
        )
        _table(doc, ["Indirect path", "β (indirect)", "95% CI", "β (direct)", "Classification"],
               [[m["path"].replace("->", "→"), _fmt(m["indirect_effect"]),
                 f"[{m['ci_95'][0]:.3f}; {m['ci_95'][1]:.3f}]" if m["ci_95"] else "",
                 _fmt(m["direct_effect"]), m["classification"]] for m in mediation])

    f2 = [s for s in assessment["structural_model"] if s["family"] == "effect_size"]
    if f2:
        doc.add_paragraph()
        doc.add_paragraph("Effect sizes f² (0.02 small / 0.15 medium / 0.35 large; Cohen, 1988):")
        _table(doc, ["Path", "f²", "Class"],
               [[s["construct"], _fmt(s["value"]), s["verdict"]] for s in f2])

    mod_f2 = [s for s in assessment["structural_model"]
              if s["family"] == "moderation_effect_size"]
    if mod_f2:
        doc.add_paragraph()
        doc.add_paragraph("Interaction-term effect sizes f² "
                          "(0.005 small / 0.01 medium / 0.025 large; Kenny, 2018):")
        _table(doc, ["Interaction", "f²", "Class"],
               [[s["construct"], _fmt(s["value"]), s["verdict"]] for s in mod_f2])

    fits = [s for s in assessment["structural_model"] if s["family"] == "model_fit"]
    if fits:
        doc.add_paragraph()
        doc.add_paragraph("Model fit: " + "; ".join(
            f"{s['metric']} = {_fmt(s['value'])} (criterion {s['threshold']}; "
            f"{s['citation']}) — {s['verdict'].upper()}" for s in fits) + ".")

    q2_blind = [s for s in assessment["structural_model"]
                if s["family"] == "predictive_relevance_q2"]
    if q2_blind:
        doc.add_paragraph()
        doc.add_paragraph(
            "Predictive relevance — blindfolding Q² (cross-validated redundancy, "
            "omission distance 7; 0.02 small / 0.15 moderate / 0.35 substantial; "
            "Hair et al., 2022):"
        )
        _table(doc, ["Endogenous construct", "Q²", "Class"],
               [[s["construct"], _fmt(s["value"]), s["verdict"]] for s in q2_blind])

    # --- Predictive power: PLSpredict ------------------------------------------
    q2 = [s for s in assessment["structural_model"] if s["family"] == "predictive_relevance"]
    power = next((s for s in assessment["structural_model"]
                  if s["family"] == "predictive_power"), None)
    if q2:
        add_section("Predictive Power (PLSpredict)")
        pred = {r["row"]: r for r in results.get("prediction") or []}
        rows = []
        for s in q2:
            r = pred.get(s["construct"], {})
            rows.append([s["construct"], _fmt(s["value"]),
                         _fmt(r.get("rmse_pls")), _fmt(r.get("rmse_lm")),
                         "PLS" if (r.get("rmse_diff") or 0) < 0 else "LM",
                         s["verdict"].upper()])
        doc.add_paragraph(
            "Out-of-sample prediction via k-fold PLSpredict (Shmueli et al., 2019). "
            "Q²predict > 0 indicates the model outperforms the indicator-mean benchmark; "
            "RMSE is compared against a linear-model (LM) benchmark:"
        )
        _table(doc, ["Indicator", "Q²predict", "RMSE (PLS)", "RMSE (LM)", "Lower RMSE", "Verdict"], rows)
        if power:
            doc.add_paragraph(
                f"PLS achieves lower out-of-sample RMSE than the LM benchmark for "
                f"{power['value']} endogenous indicators — {power['verdict']} predictive "
                f"power (Shmueli et al., 2019)."
            )

    # --- IPMA -------------------------------------------------------------------
    ipma = results.get("ipma") or {}
    performance = {r["row"]: r["performance"] for r in ipma.get("performance") or []}
    total_unstd = {r["row"]: r for r in ipma.get("total_effects_unstd") or []}
    outgoing = {p["from_construct"] for p in request["paths"]}
    endogenous = {p["to_construct"] for p in request["paths"]}
    targets = [c for c in performance if c in endogenous and c not in outgoing]
    if performance and targets:
        add_section("Importance–Performance Map Analysis (IPMA)")
        doc.add_paragraph(
            "Importance is the unstandardized total effect on the target construct; "
            "performance is the construct score rescaled to 0–100 (Ringle & Sarstedt, 2016). "
            "Constructs combining high importance with low performance are the first "
            "candidates for managerial action."
        )
        for target in targets:
            rows = []
            for pred in performance:
                if pred == target:
                    continue
                imp = total_unstd.get(pred, {}).get(target)
                if not isinstance(imp, (int, float)) or abs(imp) < 1e-9:
                    continue
                rows.append([pred, imp, performance[pred]])
            if not rows:
                continue
            rows.sort(key=lambda r: -r[1])
            doc.add_paragraph(
                f"Target: {target} (performance {_fmt(performance[target])})."
            )
            _table(doc, ["Construct", "Importance (total effect)", "Performance (0–100)"],
                   [[p, _fmt(i), f"{perf:.1f}"] for p, i, perf in rows])

    # --- Multi-group analysis ----------------------------------------------------
    if mga:
        add_section("Multi-Group Analysis (MGA)")
        m = mga["meta"]
        micom = mga["micom"]
        doc.add_paragraph(
            f"Groups compared on {m['group_variable']}: group A = {m['value_a']} "
            f"(n = {m['n_a']}), group B = {m['value_b']} (n = {m['n_b']}); "
            f"{m['effective_permutations']:,} permutations, fixed seed {m['seed']}. "
            f"Measurement invariance was tested with the MICOM procedure before any "
            f"group comparison (Henseler et al., 2016). Step 1 (configural): "
            f"{micom['step1']}."
        )
        doc.add_paragraph("Step 2 — compositional invariance (permutation test on c):")
        _table(doc, ["Construct", "c", "5% quantile", "Verdict"],
               [[s["construct"], f"{s['value']:.4f}",
                 s["threshold"].split("(")[1].rstrip(")"),
                 s["verdict"].upper()] for s in micom["step2"]])
        doc.add_paragraph()
        doc.add_paragraph("Step 3 — equality of construct means and variances (95% permutation CIs):")
        _table(doc, ["Construct", "Δ mean", "95% CI", "Equal", "Δ log variance", "95% CI", "Equal"],
               [[s["construct"], _fmt(s["mean_diff"]),
                 f"[{s['mean_ci_95'][0]:.3f}; {s['mean_ci_95'][1]:.3f}]",
                 "yes" if s["mean_equal"] else "no",
                 _fmt(s["logvar_diff"]),
                 f"[{s['var_ci_95'][0]:.3f}; {s['var_ci_95'][1]:.3f}]",
                 "yes" if s["var_equal"] else "no"] for s in micom["step3"]])
        doc.add_paragraph()
        doc.add_paragraph(
            f"MICOM verdict: {micom['invariance']} measurement invariance."
            + ("" if micom["comparison_permissible"] else f" {micom['note']}.")
        )
        if micom["comparison_permissible"]:
            doc.add_paragraph(
                "Path-coefficient differences (two-sided permutation test; "
                "Chin & Dibbern, 2010):"
            )
            _table(doc, ["Path", "β (A)", "β (B)", "Δ", "p", "Verdict"],
                   [[p["path"], _fmt(p["estimate_a"]), _fmt(p["estimate_b"]),
                     _fmt(p["difference"]), _fmt(p["p_value"]),
                     p["verdict"]] for p in mga["paths"]])

    # --- Summary ---------------------------------------------------------------
    add_section("Summary")
    summ = assessment["summary"]
    doc.add_paragraph(
        f"{summ['hypotheses_supported']} of {summ['hypotheses_total']} hypothesized paths are "
        f"supported. Measurement/structural checks: {summ['pass']} pass, {summ['review']} to "
        f"review, {summ['fail']} fail."
        + (f" {summ['indirect_effects_significant']} of {summ['indirect_effects_total']} "
           f"specific indirect effects are significant (mediation table above)."
           if summ.get("indirect_effects_total") else "")
    )
    # --- AI narrative sections (grounded in the assessment) --------------------
    if interpretation:
        sections = [
            ("Results Narrative", "results_narrative"),
            ("Discussion", "discussion"),
            ("Managerial Implications", "managerial_implications"),
            ("Limitations", "limitations"),
            ("Conclusion", "conclusion"),
        ]
        for heading, key in sections:
            text = interpretation.get(key)
            if not text:
                continue
            add_section(heading)
            for para in text.split("\n\n"):
                if para.strip():
                    doc.add_paragraph(para.strip())
        doc.add_paragraph(
            "Narrative sections were drafted by an AI writer constrained to the "
            "engine-computed statistics and threshold verdicts reported above; "
            "review before submission."
        )
    else:
        doc.add_paragraph(
            "Interpretive narrative (Discussion, Implications) has not been generated "
            "for this analysis; every statistic above is engine-computed."
        )
    return doc
