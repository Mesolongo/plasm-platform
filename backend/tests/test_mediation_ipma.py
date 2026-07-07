"""Phase 2 completion tests: mediation (specific indirect effects) and IPMA.

The extended corp-rep model doubles as the Hair et al. (2022) primer ch. 7
mediation case: COMP -> CUSA -> CUSL (0.146 x 0.505 ~ 0.074) is significant
while the direct COMP -> CUSL path (0.006) is not — indirect-only (full)
mediation; LIKE -> CUSA -> CUSL (~0.220) coexists with a significant direct
path (0.344) — complementary partial mediation. Point estimates don't depend
on bootstrap count; small nboot keeps tests fast.
"""
import io
import json
from pathlib import Path

import pandas as pd
import pytest
from fastapi.testclient import TestClient

from backend.app.main import app
from backend.app.storage import ROOT

from .helpers import create_analysis

CSV = Path(__file__).parent / "fixtures_corp_rep.csv"

client = TestClient(app)


@pytest.fixture(scope="module")
def analysis():
    with CSV.open("rb") as f:
        resp = client.post(
            "/api/datasets",
            files={"file": ("corp_rep.csv", f, "text/csv")},
            data={"missing_value": "-99"},
        )
    assert resp.status_code == 200, resp.text
    dataset = resp.json()

    spec = json.loads((ROOT / "ai" / "fixtures" / "model_spec_reference.json").read_text())
    analysis_id = create_analysis(client, {
        "dataset_id": dataset["id"], "nboot": 200, "prediction": False,
        "constructs": [{"name": c["name"], "indicators": c.get("indicators", []),
                        "measurement": c["measurement"]} for c in spec["constructs"]],
        "paths": [{"from_construct": p["from_construct"], "to_construct": p["to_construct"]}
                  for p in spec["paths"]],
    })["id"]
    return {
        "id": analysis_id,
        "results": client.get(f"/api/analyses/{analysis_id}/results").json(),
        "assessment": client.get(f"/api/analyses/{analysis_id}/assessment").json(),
    }


def test_specific_indirect_matches_published_values(analysis):
    rows = {r["row"]: r for r in analysis["results"]["specific_indirect"]}
    # Published ch. 7 values: 0.146 x 0.505 and 0.436 x 0.505
    assert abs(rows["COMP->CUSA->CUSL"]["Original Est."] - 0.074) < 0.005
    assert abs(rows["LIKE->CUSA->CUSL"]["Original Est."] - 0.220) < 0.005
    # Serial chains from the exogenous drivers are enumerated too
    assert "QUAL->COMP->CUSA->CUSL" in rows
    assert all(r["2.5% CI"] is not None for r in rows.values())


def test_mediation_classification(analysis):
    med = {m["path"]: m for m in analysis["assessment"]["mediation"]}
    comp = med["COMP -> CUSA -> CUSL"]
    assert comp["significant"] and comp["classification"] == "indirect-only (full mediation)"
    like = med["LIKE -> CUSA -> CUSL"]
    assert like["significant"] and like["classification"] == "complementary (partial mediation)"
    assert like["direct_effect"] == pytest.approx(0.344, abs=0.005)
    # No direct QUAL -> CUSL path is modeled: typed as indirect-only via the drivers
    qual = med["QUAL -> COMP -> CUSA -> CUSL"]
    assert "no direct path modeled" in qual["classification"]
    summary = analysis["assessment"]["summary"]
    assert summary["indirect_effects_total"] == len(med) > 0
    assert summary["indirect_effects_significant"] >= 2


def test_ipma(analysis):
    ipma = analysis["results"]["ipma"]
    perf = {r["row"]: r["performance"] for r in ipma["performance"]}
    assert all(0 <= v <= 100 for v in perf.values())

    # Single-item construct: performance is the mean-replaced item rescaled 0-100
    s = pd.read_csv(CSV)["cusa"].astype(float).replace(-99.0, float("nan"))
    s = s.fillna(s.mean())
    expected = ((s - s.min()) / (s.max() - s.min()) * 100).mean()
    assert perf["CUSA"] == pytest.approx(expected, abs=1e-3)

    # Importance ordering on the final outcome mirrors the standardized totals:
    # LIKE (0.344 + 0.220) far outweighs COMP (0.006 + 0.074)
    te = {r["row"]: r for r in ipma["total_effects_unstd"]}
    assert te["LIKE"]["CUSL"] > te["COMP"]["CUSL"] > 0
    assert te["QUAL"]["CUSL"] > 0  # drivers reach the outcome through the chains


def test_report_includes_mediation_and_ipma(analysis):
    resp = client.get(f"/api/analyses/{analysis['id']}/report.docx")
    assert resp.status_code == 200
    from docx import Document
    doc = Document(io.BytesIO(resp.content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Zhao" in text and "Importance–Performance Map Analysis" in text
    cells = {c.text for t in doc.tables for row in t.rows for c in row.cells}
    assert "COMP → CUSA → CUSL" in cells
    assert "indirect-only (full mediation)" in cells
